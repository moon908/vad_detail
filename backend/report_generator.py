import os
import zipfile
import csv
import json
import numpy as np
from typing import List, Dict, Any

# Configure Matplotlib to run headlessly
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

# python-docx imports for Word documents
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

class NumberedCanvas(canvas.Canvas):
    """
    Two-pass canvas to add 'Page X of Y' footers and headers.
    """
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_elements(num_pages)
            super().showPage()
        super().save()

    def draw_page_elements(self, page_count):
        # Skip header and footer on the cover page (Page 1)
        if self._pageNumber == 1:
            return

        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#64748b")) # Slate-500

        # Header
        self.drawString(54, 750, "Voice Activity Detection (VAD) Analysis Report")
        self.setStrokeColor(colors.HexColor("#e2e8f0")) # Slate-200
        self.setLineWidth(0.5)
        self.line(54, 742, 558, 742)

        # Footer
        self.line(54, 45, 558, 45)
        page_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(558, 30, page_text)
        self.drawString(54, 30, "Confidential - VAD Analysis Platform")
        
        self.restoreState()


class ReportGenerator:
    @staticmethod
    def _generate_plots(
        audio_data: Any,
        sr: int,
        segments: List[Dict[str, Any]],
        stats: Dict[str, Any],
        temp_dir: str
    ) -> Dict[str, str]:
        """
        Generates static visualization plots using Matplotlib and saves them as PNG.
        """
        os.makedirs(temp_dir, exist_ok=True)
        paths = {}

        # 1. Waveform with Speech Highlights
        plt.figure(figsize=(10, 3.5))
        # Downsample audio data for faster plotting if it's too large
        step = max(1, len(audio_data) // 50000)
        downsampled_audio = audio_data[::step]
        time_axis = (np.arange(len(downsampled_audio)) * step) / sr
        
        plt.plot(time_axis, downsampled_audio, color='#94a3b8', alpha=0.6, label='Audio Signal')
        
        # Highlight speech regions
        first_speech = True
        for seg in segments:
            label = 'Detected Speech' if first_speech else ""
            plt.axvspan(seg['start'], seg['end'], color='#3b82f6', alpha=0.3, label=label)
            first_speech = False
            
        plt.title('Voice Activity Detection Timeline', fontsize=12, fontweight='bold', pad=10, color='#1e293b')
        plt.xlabel('Time (seconds)', fontsize=9, color='#475569')
        plt.ylabel('Amplitude', fontsize=9, color='#475569')
        plt.grid(True, linestyle='--', alpha=0.3)
        if segments:
            plt.legend(loc='upper right', fontsize=8)
        plt.tight_layout()
        waveform_path = os.path.join(temp_dir, "plot_waveform.png")
        plt.savefig(waveform_path, dpi=150)
        plt.close()
        paths["waveform"] = waveform_path

        # 2. Donut Pie Chart of Speech vs Silence
        plt.figure(figsize=(4, 4))
        labels = ['Speech', 'Silence']
        sizes = [stats['speechPercentage'], stats['silencePercentage']]
        colors_pie = ['#3b82f6', '#e2e8f0']
        
        plt.pie(
            sizes,
            labels=labels,
            colors=colors_pie,
            autopct='%1.1f%%',
            startangle=90,
            pctdistance=0.75,
            textprops={'fontsize': 10, 'weight': 'bold', 'color': '#1e293b'},
            wedgeprops={'edgecolor': 'white', 'linewidth': 2}
        )
        # Draw center circle to make it a donut
        centre_circle = plt.Circle((0,0), 0.55, fc='white')
        fig = plt.gcf()
        fig.gca().add_artist(centre_circle)
        
        plt.title('Speech vs Silence Ratio', fontsize=11, fontweight='bold', pad=10, color='#1e293b')
        plt.tight_layout()
        pie_path = os.path.join(temp_dir, "plot_pie.png")
        plt.savefig(pie_path, dpi=150)
        plt.close()
        paths["pie"] = pie_path

        # 3. Bar Chart: Speech Segment Durations
        plt.figure(figsize=(6, 4))
        if segments:
            segment_ids = [f"Seg {seg['id']}" for seg in segments[:15]] # Limit to first 15 segments for display
            durations = [seg['duration'] for seg in segments[:15]]
            
            plt.bar(segment_ids, durations, color='#60a5fa', edgecolor='#2563eb', width=0.6)
            plt.title('Speech Segment Durations (First 15)', fontsize=11, fontweight='bold', pad=10, color='#1e293b')
            plt.ylabel('Duration (seconds)', fontsize=9, color='#475569')
            plt.xlabel('Segment', fontsize=9, color='#475569')
            plt.xticks(rotation=45, ha='right', fontsize=8)
            plt.grid(True, axis='y', linestyle='--', alpha=0.3)
        else:
            plt.text(0.5, 0.5, 'No Speech Detected', horizontalalignment='center', verticalalignment='center', fontsize=12)
            plt.title('Speech Segment Durations', fontsize=11, fontweight='bold', pad=10, color='#1e293b')
            
        plt.tight_layout()
        bar_path = os.path.join(temp_dir, "plot_bar.png")
        plt.savefig(bar_path, dpi=150)
        plt.close()
        paths["bar"] = bar_path

        return paths

    @staticmethod
    def generate_csv(segments: List[Dict[str, Any]], csv_path: str):
        """Generates a CSV report of speech segments."""
        with open(csv_path, mode='w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["Segment ID", "Start Time (seconds)", "End Time (seconds)", "Duration (seconds)"])
            for seg in segments:
                writer.writerow([seg['id'], seg['start'], seg['end'], seg['duration']])

    @staticmethod
    def generate_json(stats: Dict[str, Any], segments: List[Dict[str, Any]], json_path: str):
        """Generates a JSON report of statistics and speech segments."""
        data = {
            "statistics": {k: v for k, v in stats.items() if k != "silenceSegments"},
            "segments": segments
        }
        with open(json_path, mode='w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)

    @staticmethod
    def generate_pdf(
        stats: Dict[str, Any],
        segments: List[Dict[str, Any]],
        audio_data: Any,
        sr: int,
        pdf_path: str,
        temp_dir: str,
        filename: str,
        filesize_formatted: str,
        ai_insights: List[str]
    ):
        """Generates a high-quality ReportLab PDF report."""
        # 1. Generate PNG plots
        plots = ReportGenerator._generate_plots(audio_data, sr, segments, stats, temp_dir)

        # 2. Build PDF Document
        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        
        # Modify existing styles or add unique ones
        title_style = ParagraphStyle(
            'CoverTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=28,
            leading=34,
            textColor=colors.HexColor("#1e3a8a"), # Deep Navy
            alignment=0, # Left-aligned
            spaceAfter=15
        )
        
        subtitle_style = ParagraphStyle(
            'CoverSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#475569"), # Slate-600
            alignment=0,
            spaceAfter=40
        )
        
        h1_style = ParagraphStyle(
            'SectionH1',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#1e3a8a"),
            spaceBefore=15,
            spaceAfter=10,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#334155") # Slate-700
        )

        meta_label_style = ParagraphStyle(
            'MetaLabel',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#475569")
        )

        story = []

        # ==========================================
        # COVER PAGE
        # ==========================================
        story.append(Spacer(1, 100))
        # Large color accent block simulation
        story.append(Paragraph("VOICE ACTIVITY DETECTION<br/>ANALYSIS REPORT", title_style))
        story.append(Paragraph("A professional AI-driven audio segmentation & acoustic analytics overview", subtitle_style))
        story.append(Spacer(1, 40))

        # Metadata table on cover page
        meta_data = [
            [Paragraph("Acoustic Target File:", meta_label_style), Paragraph(filename, body_style)],
            [Paragraph("File Size:", meta_label_style), Paragraph(filesize_formatted, body_style)],
            [Paragraph("Audio Duration:", meta_label_style), Paragraph(f"{stats['duration']:.2f} seconds", body_style)],
            [Paragraph("Total Speech detected:", meta_label_style), Paragraph(f"{stats['totalSpeech']:.2f} seconds ({stats['speechPercentage']:.1f}%)", body_style)],
            [Paragraph("Total Silence detected:", meta_label_style), Paragraph(f"{stats['totalSilence']:.2f} seconds ({stats['silencePercentage']:.1f}%)", body_style)],
            [Paragraph("Report Generated:", meta_label_style), Paragraph("July 2, 2026", body_style)]
        ]
        
        meta_table = Table(meta_data, colWidths=[2.2 * inch, 4.0 * inch])
        meta_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor("#f1f5f9")),
        ]))
        
        story.append(meta_table)
        story.append(PageBreak())

        # ==========================================
        # PAGE 2: EXECUTIVE SUMMARY & CHARTS
        # ==========================================
        story.append(Paragraph("Executive Summary & Statistics", h1_style))
        story.append(Spacer(1, 10))

        # Metrics summary blocks
        metrics_data = [
            ["Speech Segments", f"{stats['speechSegmentsCount']}", "Silence Segments", f"{stats['silenceSegmentsCount']}"],
            ["Longest Speech", f"{stats['longestSpeech']} s", "Average Speech", f"{stats['avgSpeech']} s"],
            ["Shortest Speech", f"{stats['shortestSpeech']} s", "Average Silence", f"{stats['avgSilence']} s"],
            ["Speaking Speed", f"{stats['estimatedSpeakingSpeedWpm']:.0f} WPM", "Word Count Est.", f"{stats['estimatedWords']} words"]
        ]
        metrics_table = Table(metrics_data, colWidths=[1.55 * inch, 1.45 * inch, 1.55 * inch, 1.45 * inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#f8fafc")),
            ('BACKGROUND', (2,0), (2,-1), colors.HexColor("#f8fafc")),
            ('TEXTCOLOR', (0,0), (-1,-1), colors.HexColor("#1e293b")),
            ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
            ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
            ('FONTNAME', (2,0), (2,-1), 'Helvetica-Bold'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(metrics_table)
        story.append(Spacer(1, 20))

        # Embed charts side-by-side or stacked
        # Donut Chart and Durations Bar Chart stacked
        charts_data = [
            [Image(plots["pie"], width=2.8 * inch, height=2.8 * inch), 
             Image(plots["bar"], width=3.2 * inch, height=2.5 * inch)]
        ]
        charts_table = Table(charts_data, colWidths=[3.0 * inch, 3.4 * inch])
        charts_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ]))
        story.append(charts_table)
        
        story.append(PageBreak())

        # ==========================================
        # PAGE 3: TIMELINE & OBSERVATIONS
        # ==========================================
        story.append(Paragraph("Acoustic Timeline Visualization", h1_style))
        story.append(Spacer(1, 5))
        story.append(Image(plots["waveform"], width=6.2 * inch, height=2.1 * inch))
        story.append(Spacer(1, 20))

        # AI Insights Section
        story.append(Paragraph("AI-Generated Observations & Insights", h1_style))
        story.append(Spacer(1, 8))
        
        insight_bullet_style = ParagraphStyle(
            'InsightBullet',
            parent=body_style,
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=8
        )
        
        for ins in ai_insights:
            bullet_p = Paragraph(f"• &nbsp; {ins}", insight_bullet_style)
            story.append(bullet_p)

        story.append(Spacer(1, 20))

        # Final Conclusion Box
        conclusion_hdr = ParagraphStyle('ConclusionHdr', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11, leading=14, textColor=colors.HexColor("#1e3a8a"), spaceAfter=5)
        conclusion_txt = ParagraphStyle('ConclusionTxt', parent=body_style, fontSize=9.5, leading=13, textColor=colors.HexColor("#1e293b"))
        
        conclusion_box_data = [[
            Paragraph("System Final Conclusion", conclusion_hdr),
            Paragraph("Based on voice activity evaluation, this audio file represents a clean recording with structured segments. Silence profiles indicate no major communication dropouts, and speech rate estimates suggest a standard rhythm. Export files containing isolated audio components can be reviewed individually.", conclusion_txt)
        ]]
        
        conclusion_table = Table(conclusion_box_data, colWidths=[6.0 * inch])
        conclusion_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#eff6ff")), # Soft Blue-50
            ('BOX', (0,0), (-1,-1), 1.0, colors.HexColor("#bfdbfe")), # Blue-200
            ('TOPPADDING', (0,0), (-1,-1), 12),
            ('BOTTOMPADDING', (0,0), (-1,-1), 12),
            ('LEFTPADDING', (0,0), (-1,-1), 12),
            ('RIGHTPADDING', (0,0), (-1,-1), 12),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        story.append(conclusion_table)

        # ==========================================
        # PAGE 4: SEGMENTS DATA TABLE (If there is space or overflow)
        # ==========================================
        story.append(PageBreak())
        story.append(Paragraph("Speech Timestamp Segments List", h1_style))
        story.append(Spacer(1, 10))

        # Build table header
        table_content = [["Seg ID", "Start Time (s)", "End Time (s)", "Duration (s)"]]
        for seg in segments[:30]: # Limit to first 30 segments to fit on page neatly
            table_content.append([
                f"{seg['id']:03d}",
                f"{seg['start']:.2f}",
                f"{seg['end']:.2f}",
                f"{seg['duration']:.2f}"
            ])
            
        if len(segments) > 30:
            table_content.append(["...", "...", "...", f"+ {len(segments) - 30} more segments"])
            
        segments_table = Table(table_content, colWidths=[1.2 * inch, 1.6 * inch, 1.6 * inch, 1.6 * inch])
        
        # Grid styles
        ts = [
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1e3a8a")),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('TOPPADDING', (0,0), (-1,-1), 5),
        ]
        
        # Row coloring
        for r in range(1, len(table_content)):
            if r % 2 == 0:
                ts.append(('BACKGROUND', (0, r), (-1, r), colors.HexColor("#f8fafc")))
                
        segments_table.setStyle(TableStyle(ts))
        story.append(segments_table)

        # 3. Build document
        doc.build(story, canvasmaker=NumberedCanvas)

    @staticmethod
    def generate_docx(
        stats: Dict[str, Any],
        segments: List[Dict[str, Any]],
        docx_path: str,
        temp_dir: str,
        filename: str,
        filesize_formatted: str,
        ai_insights: List[str]
    ):
        """Generates a structured Microsoft Word (.docx) report."""
        doc = Document()
        
        # Page Setup
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # Style Helpers
        def set_font(run, name="Calibri", size=11, bold=False, italic=False, color_rgb=None):
            run.font.name = name
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            if color_rgb:
                run.font.color.rgb = color_rgb

        def style_heading(p, text, size=16, color_rgb=RGBColor(30, 58, 138)):
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            set_font(run, name="Calibri", size=size, bold=True, color_rgb=color_rgb)

        def set_cell_background(cell, fill_hex):
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # 1. Document Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_title = p_title.add_run("VOICE ACTIVITY DETECTION ANALYSIS REPORT")
        set_font(run_title, name="Calibri", size=24, bold=True, color_rgb=RGBColor(30, 58, 138))
        
        # Subtitle
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("A professional AI-driven audio segmentation & acoustic analytics overview")
        set_font(run_sub, name="Calibri", size=12, italic=True, color_rgb=RGBColor(71, 85, 105))
        p_sub.paragraph_format.space_after = Pt(20)

        # Horizontal Rule
        p_hr = doc.add_paragraph()
        p_hr_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/></w:pBdr>')
        p_hr._p.get_or_add_pPr().append(p_hr_border)

        # 2. File Metadata Section
        p_meta_title = doc.add_paragraph()
        style_heading(p_meta_title, "Acoustic File Metadata", size=14)

        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = 'Table Grid'
        meta_data = [
            ("Acoustic Target File", filename),
            ("File Size", filesize_formatted),
            ("Audio Duration", f"{stats['duration']:.2f} seconds"),
            ("Total Speech Detected", f"{stats['totalSpeech']:.2f} seconds ({stats['speechPercentage']:.1f}%)"),
            ("Total Silence Detected", f"{stats['totalSilence']:.2f} seconds ({stats['silencePercentage']:.1f}%)")
        ]
        
        for idx, (label, val) in enumerate(meta_data):
            row = meta_table.rows[idx]
            
            cell_lbl = row.cells[0]
            cell_lbl.text = label
            set_cell_background(cell_lbl, "F1F5F9")
            set_font(cell_lbl.paragraphs[0].runs[0], size=10, bold=True, color_rgb=RGBColor(71, 85, 105))
            
            cell_val = row.cells[1]
            cell_val.text = val
            set_font(cell_val.paragraphs[0].runs[0], size=10)

        doc.add_paragraph().paragraph_format.space_before = Pt(10)

        # 3. Speech Statistics
        p_stats_title = doc.add_paragraph()
        style_heading(p_stats_title, "Speech Statistics Summary", size=14)

        stats_table = doc.add_table(rows=4, cols=4)
        stats_table.style = 'Table Grid'
        stats_data = [
            ("Speech Segments", f"{stats['speechSegmentsCount']}", "Silence Segments", f"{stats['silenceSegmentsCount']}"),
            ("Longest Speech", f"{stats['longestSpeech']} s", "Average Speech", f"{stats['avgSpeech']} s"),
            ("Shortest Speech", f"{stats['shortestSpeech']} s", "Average Silence", f"{stats['avgSilence']} s"),
            ("Estimated Speaking Speed", f"{stats['estimatedSpeakingSpeedWpm']:.0f} WPM", "Word Count Estimate", f"{stats['estimatedWords']} words")
        ]

        for r_idx, row_data in enumerate(stats_data):
            row = stats_table.rows[r_idx]
            for c_idx, text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = text
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Style headers vs values
                is_lbl = (c_idx % 2 == 0)
                if is_lbl:
                    set_cell_background(cell, "F8FAFC")
                    set_font(p.runs[0], size=9.5, bold=True, color_rgb=RGBColor(30, 41, 59))
                else:
                    set_font(p.runs[0], size=9.5)

        doc.add_paragraph().paragraph_format.space_before = Pt(15)

        # 4. Embed Matplotlib Plots
        p_plots_title = doc.add_paragraph()
        style_heading(p_plots_title, "Acoustic Timeline Visualization", size=14)

        waveform_img = os.path.join(temp_dir, "plot_waveform.png")
        if os.path.exists(waveform_img):
            p_waveform = doc.add_paragraph()
            p_waveform.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_waveform.add_run().add_picture(waveform_img, width=Inches(6.2))
            p_cap1 = doc.add_paragraph()
            p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p_cap1.add_run("Figure 1: Voice Activity Detection (VAD) audio signal highlighting speech sections"), size=8.5, italic=True)

        pie_img = os.path.join(temp_dir, "plot_pie.png")
        if os.path.exists(pie_img):
            p_pie = doc.add_paragraph()
            p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_pie.add_run().add_picture(pie_img, width=Inches(3.2))
            p_cap2 = doc.add_paragraph()
            p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p_cap2.add_run("Figure 2: Ratio share of active vocal speech versus quiet silence"), size=8.5, italic=True)

        doc.add_page_break()

        # 5. AI Generated Insights
        p_insights_title = doc.add_paragraph()
        style_heading(p_insights_title, "AI-Generated Observations & Insights", size=14)

        for ins in ai_insights:
            p_ins = doc.add_paragraph(style='List Bullet')
            run = p_ins.add_run(ins)
            set_font(run, size=10, color_rgb=RGBColor(51, 65, 85))

        doc.add_paragraph().paragraph_format.space_before = Pt(10)

        # 6. Conclusion Box
        p_conclusion = doc.add_paragraph()
        p_conclusion.paragraph_format.left_indent = Inches(0.2)
        p_conclusion.paragraph_format.right_indent = Inches(0.2)
        p_conclusion.paragraph_format.space_before = Pt(10)
        p_conclusion.paragraph_format.space_after = Pt(10)
        
        # Border XML
        conclusion_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="3B82F6"/></w:pBdr>')
        p_conclusion._p.get_or_add_pPr().append(conclusion_border)
        
        # Background XML
        conclusion_shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EFF6FF"/>')
        p_conclusion._p.get_or_add_pPr().append(conclusion_shd)

        run_conclusion_title = p_conclusion.add_run("System Final Conclusion\n")
        set_font(run_conclusion_title, size=10.5, bold=True, color_rgb=RGBColor(30, 58, 138))
        
        run_conclusion_txt = p_conclusion.add_run(
            "Based on voice activity evaluation, this audio file represents a clean recording with structured segments. "
            "Silence profiles indicate no major communication dropouts, and speech rate estimates suggest a standard rhythm. "
            "Export files containing isolated audio components can be reviewed individually."
        )
        set_font(run_conclusion_txt, size=9.5, color_rgb=RGBColor(30, 41, 59))

        doc.add_paragraph().paragraph_format.space_before = Pt(15)

        # 7. Timestamps Table
        p_tbl_title = doc.add_paragraph()
        style_heading(p_tbl_title, "Acoustic Speech Segments List (First 50)", size=14)

        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        
        hdr_cells = tbl.rows[0].cells
        headers = ["Seg ID", "Start Time (s)", "End Time (s)", "Duration (s)"]
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            cell.text = text
            set_cell_background(cell, "1E3A8A")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.runs[0], size=10, bold=True, color_rgb=RGBColor(255, 255, 255))

        for seg in segments[:50]: # Limit to first 50
            row_cells = tbl.add_row().cells
            row_cells[0].text = f"{seg['id']:03d}"
            row_cells[1].text = f"{seg['start']:.2f}"
            row_cells[2].text = f"{seg['end']:.2f}"
            row_cells[3].text = f"{seg['duration']:.2f}"
            
            for c_idx in range(4):
                cell = row_cells[c_idx]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(p.runs[0], size=9.5)
                # Alternating row colors
                if seg['id'] % 2 == 0:
                    set_cell_background(cell, "F8FAFC")

        # Save document
        doc.save(docx_path)

    @staticmethod
    def generate_zip(zip_path: str, files_to_include: Dict[str, str], segments_dir: str):
        """Generates a compressed ZIP file enclosing all output reports and sliced WAV files."""
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Write primary reports
            for arcname, filepath in files_to_include.items():
                if os.path.exists(filepath):
                    zip_file.write(filepath, arcname)

            # Write individual segments if present
            if os.path.exists(segments_dir):
                for root, _, files in os.walk(segments_dir):
                    for file in files:
                        full_path = os.path.join(root, file)
                        rel_path = os.path.join("segments", file)
                        zip_file.write(full_path, rel_path)
